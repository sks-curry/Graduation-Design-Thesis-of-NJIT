from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


LATEX_SPECIALS = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def iter_block_items(parent: DocumentObject) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def latex_escape(text: str) -> str:
    return "".join(LATEX_SPECIALS.get(ch, ch) for ch in text)


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def paragraph_image_names(paragraph: Paragraph) -> list[str]:
    names: list[str] = []
    for blip in paragraph._element.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        rel = paragraph.part.rels.get(rel_id)
        if rel is None:
            continue
        names.append(Path(rel.target_ref).name)
    return names


def extract_media(docx_path: Path, assets_dir: Path) -> list[Path]:
    assets_dir.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            target = assets_dir / Path(name).name
            with zf.open(name) as src, target.open("wb") as dst:
                shutil.copyfileobj(src, dst)
            extracted.append(target)
    return sorted(extracted, key=lambda p: p.name)


def table_to_latex(table: Table) -> list[str]:
    rows = []
    max_cols = max((len(row.cells) for row in table.rows), default=0)
    if max_cols == 0:
        return []
    if max_cols == 2:
        colspec = r">{\centering\arraybackslash}p{0.28\textwidth}|X"
    else:
        colspec = "|".join(["X"] * max_cols)
    rows.append(r"\begin{table}[htbp]")
    rows.append(r"\centering")
    rows.append(r"\zihao{5}")
    rows.append(r"\renewcommand{\arraystretch}{1.25}")
    rows.append(r"\begin{tabularx}{0.88\textwidth}{" + "|" + colspec + "|" + "}")
    rows.append(r"\hline")
    for row in table.rows:
        cells = [latex_escape(normalize_text(cell.text)) for cell in row.cells]
        cells += [""] * (max_cols - len(cells))
        rows.append(" & ".join(cells) + r" \\")
        rows.append(r"\hline")
    rows.append(r"\end{tabularx}")
    rows.append(r"\end{table}")
    return rows


def graphics_block(image_path: Path) -> list[str]:
    rel_path = image_path.as_posix()
    if image_path.suffix.lower() in {".emf", ".wmf"}:
        return [
            rf"% Unsupported Word vector image omitted: {rel_path}",
            r"\begin{center}",
            rf"\fbox{{\parbox{{0.75\textwidth}}{{Word vector image placeholder: {latex_escape(image_path.name)}}}}}",
            r"\end{center}",
        ]
    return [
        r"\begin{figure}[htbp]",
        r"\centering",
        rf"\includegraphics[width=0.82\textwidth]{{{rel_path}}}",
        r"\end{figure}",
    ]


def paragraph_to_latex(paragraph: Paragraph, asset_lookup: dict[str, Path]) -> list[str]:
    style = paragraph.style.name if paragraph.style is not None else ""
    text = normalize_text(paragraph.text)
    lines: list[str] = []

    if style.startswith("toc ") or text in {"目 录", "目    录"}:
        return []

    for image_name in paragraph_image_names(paragraph):
        image = asset_lookup.get(image_name)
        if image is not None:
            lines.extend(graphics_block(image))

    if not text:
        return lines

    escaped = latex_escape(text)
    if style == "Heading 1":
        lines.append(rf"\chapter{{{escaped}}}")
    elif style == "Heading 2":
        lines.append(rf"\section{{{escaped}}}")
    elif style == "Heading 3":
        lines.append(rf"\subsection{{{escaped}}}")
    elif style == "Heading 4":
        lines.append(rf"\subsubsection{{{escaped}}}")
    elif re.match(r"^图\d+(\.\d+)?\s+", text):
        lines.append(rf"\begin{{center}}\zihao{{5}}{escaped}\end{{center}}")
    elif re.match(r"^表\d+(\.\d+)?\s+", text):
        lines.append(rf"\begin{{center}}\zihao{{5}}{escaped}\end{{center}}")
    elif text.startswith("关键词"):
        lines.append(rf"\noindent\textbf{{{escaped}}}")
    elif text.startswith("Key words") or text.startswith("Keywords"):
        lines.append(rf"\noindent\textbf{{{escaped}}}")
    elif text in {"摘 要", "摘       要"}:
        lines.append(r"\chapter*{摘\quad 要}")
        lines.append(r"\addcontentsline{toc}{chapter}{摘\quad 要}")
    elif text == "ABSTRACT":
        lines.append(r"\chapter*{ABSTRACT}")
        lines.append(r"\addcontentsline{toc}{chapter}{ABSTRACT}")
    elif re.match(r"^（\d+）", text):
        lines.append(rf"\noindent {escaped}\\")
    else:
        lines.append(escaped)
    return lines


def collect_plain_paragraphs(document: DocumentObject) -> list[str]:
    return [normalize_text(p.text) for p in document.paragraphs if normalize_text(p.text)]


def table_cells(table: Table) -> list[list[str]]:
    return [[normalize_text(cell.text) for cell in row.cells] for row in table.rows]


def front_pages(document: DocumentObject) -> list[str]:
    all_text = collect_plain_paragraphs(document)
    zh_title = "双电机同轴驱动控制系统设计"
    meta = {
        "姓   名": "学生姓名",
        "学   号": "203220104",
        "班   级": "自动化221",
        "学   院": "自动化学院",
        "专   业": "自动化",
        "指导教师": "姓名1  姓名2",
        "日期": "2026年5月",
    }
    if len(document.tables) >= 1:
        first = table_cells(document.tables[0])
        if first and len(first[0]) > 1:
            zh_title = first[0][1]
    if len(document.tables) >= 2:
        for row in table_cells(document.tables[1]):
            if len(row) < 2:
                continue
            key = row[0].replace("：", "").strip()
            value = row[1].replace("\n", " ").strip()
            if key:
                meta[key] = value
            elif value:
                meta["日期"] = value

    english = []
    for text in all_text:
        if text in {"摘 要", "摘       要"}:
            break
        if re.search(r"[A-Za-z]", text):
            english.append(text)

    title_width = "9.5cm"
    lines = [
        r"\begin{titlepage}",
        r"\thispagestyle{empty}",
        r"\centering",
        r"\vspace*{1.2cm}",
        r"{\zihao{-2}\heiti 本科生毕业论文（设计）\par}",
        r"\vspace{2.2cm}",
        r"\renewcommand{\arraystretch}{1.8}",
        r"\begin{tabular}{rl}",
        rf"\zihao{{4}}\heiti 题目： & \underline{{\makebox[{title_width}][c]{{\zihao{{4}} {latex_escape(zh_title)}}}}}\\",
        r"\end{tabular}",
        r"\vspace{2.0cm}",
        r"\begin{tabular}{rl}",
    ]
    for key in ["姓   名", "学   号", "班   级", "学   院", "专   业", "指导教师"]:
        value = meta.get(key, "")
        lines.append(
            rf"\zihao{{4}}\heiti {latex_escape(key)}： & \underline{{\makebox[{title_width}][c]{{\zihao{{4}} {latex_escape(value)}}}}}\\"
        )
    lines.extend(
        [
            r"\end{tabular}",
            r"\vfill",
            rf"{{\zihao{{4}} {latex_escape(meta.get('日期', '2026年5月'))}\par}}",
            r"\end{titlepage}",
            r"\clearpage",
        ]
    )

    if english:
        lines.extend(
            [
                r"\begin{titlepage}",
                r"\thispagestyle{empty}",
                r"\centering",
                r"\vspace*{1.0cm}",
            ]
        )
        for i, text in enumerate(english):
            if i == 0:
                lines.append(rf"{{\zihao{{3}}\bfseries {latex_escape(text)}\par}}")
                lines.append(r"\vspace{1.4cm}")
            elif i == 1:
                lines.append(rf"{{\zihao{{3}}\bfseries {latex_escape(text)}\par}}")
                lines.append(r"\vspace{1.2cm}")
            else:
                lines.append(rf"{{\zihao{{4}} {latex_escape(text)}\par}}")
                lines.append(r"\vspace{0.35cm}")
        lines.extend([r"\end{titlepage}", r"\clearpage"])
    return lines


def build_latex(docx_path: Path, output_tex: Path, assets_dir: Path) -> None:
    if assets_dir.exists():
        shutil.rmtree(assets_dir)
    document = Document(docx_path)
    media = extract_media(docx_path, assets_dir)
    asset_lookup = {path.name: Path("latex_assets") / path.name for path in media}

    preamble = [
        r"\documentclass[UTF8,zihao=-4,openany,oneside]{ctexbook}",
        r"\usepackage[a4paper,left=3cm,right=2.6cm,top=3.5cm,bottom=2.6cm,headsep=0.7cm]{geometry}",
        r"\usepackage{fontspec}",
        r"\usepackage{graphicx}",
        r"\usepackage{caption}",
        r"\usepackage{booktabs}",
        r"\usepackage{array}",
        r"\usepackage{tabularx}",
        r"\usepackage{longtable}",
        r"\usepackage{setspace}",
        r"\usepackage{fancyhdr}",
        r"\usepackage{titlesec}",
        r"\usepackage{tocloft}",
        r"\usepackage{indentfirst}",
        r"\usepackage{hyperref}",
        r"\hypersetup{hidelinks}",
        r"\setmainfont{Times New Roman}",
        r"\setCJKmainfont{SimSun}",
        r"\setCJKsansfont{SimHei}",
        r"\setCJKfamilyfont{hei}{SimHei}",
        r"\providecommand{\heiti}{}",
        r"\renewcommand{\heiti}{\CJKfamily{hei}}",
        r"\setlength{\parindent}{2em}",
        r"\setlength{\parskip}{0pt}",
        r"\linespread{1.67}",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        r"\fancyhead[C]{\zihao{5} 南京工程学院自动化学院本科毕业设计（论文）}",
        r"\fancyfoot[C]{\thepage}",
        r"\renewcommand{\headrulewidth}{0.4pt}",
        r"\ctexset{",
        r"  chapter={format=\centering\zihao{3}\heiti,beforeskip=0pt,afterskip=20pt,name={第,章},number=\chinese{chapter}},",
        r"  section={format=\zihao{4}\heiti,beforeskip=12pt,afterskip=6pt},",
        r"  subsection={format=\zihao{-4}\heiti,beforeskip=6pt,afterskip=6pt}",
        r"}",
        r"\captionsetup{font=small,labelsep=space}",
        r"\renewcommand{\contentsname}{目\quad 录}",
        r"\renewcommand{\cfttoctitlefont}{\hfill\zihao{3}\heiti}",
        r"\renewcommand{\cftaftertoctitle}{\hfill}",
        r"\renewcommand{\cftchapfont}{\zihao{-4}}",
        r"\renewcommand{\cftsecfont}{\zihao{-4}}",
        r"\renewcommand{\cftchappagefont}{\zihao{-4}}",
        r"\renewcommand{\cftsecpagefont}{\zihao{-4}}",
        r"\begin{document}",
    ]

    body: list[str] = front_pages(document)
    body.extend(
        [
            r"\frontmatter",
            r"\pagenumbering{Roman}",
            r"\setcounter{page}{1}",
        ]
    )
    in_main = False
    seen_content = False
    skipped_cover_tables = 0
    for block in iter_block_items(document):
        if isinstance(block, Table) and skipped_cover_tables < 2:
            skipped_cover_tables += 1
            continue
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            style = block.style.name if block.style is not None else ""
            if style.startswith("toc ") or text in {"目 录", "目    录"}:
                continue
            if text == "Graduation Design (Thesis)" or (not seen_content and text.startswith("本科生毕业论文")):
                continue
            if text in {"摘 要", "摘       要"}:
                seen_content = True
            if style == "Heading 1" and not in_main:
                body.extend(
                    [
                        r"\clearpage",
                        r"\tableofcontents",
                        r"\clearpage",
                        r"\mainmatter",
                        r"\pagenumbering{arabic}",
                        r"\setcounter{page}{1}",
                    ]
                )
                in_main = True
            body.extend(paragraph_to_latex(block, asset_lookup))
            if body and body[-1]:
                body.append("")
        else:
            body.extend(table_to_latex(block))
            body.append("")

    ending = [r"\end{document}", ""]
    output_tex.parent.mkdir(parents=True, exist_ok=True)
    output_tex.write_text("\n".join(preamble + body + ending), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--assets-dir", type=Path, required=True)
    args = parser.parse_args()
    build_latex(args.docx, args.out, args.assets_dir)


if __name__ == "__main__":
    main()
